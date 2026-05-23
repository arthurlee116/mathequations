import tempfile
import unittest
from pathlib import Path

from docx import Document
from PIL import Image

from render_docx_functions import render_docx_to_image


class RenderDocxFunctionsScriptTests(unittest.TestCase):
    def test_render_docx_to_image_fills_polygons_from_docx_functions(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            docx_path = root / "functions.docx"
            output_path = root / "render.png"

            doc = Document()
            doc.add_heading("1. red", level=2)
            doc.add_paragraph("颜色：#ff0000 | 函数数量：4")
            table = doc.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "序号"
            table.rows[0].cells[1].text = "函数"
            rows = [
                "一次函数：y = 0x + 10（定义域：0 <= x <= 10）",
                "竖直直线：x = 10（y 的范围：0 <= y <= 10）",
                "一次函数：y = 0x + 0（定义域：0 <= x <= 10）",
                "竖直直线：x = 0（y 的范围：0 <= y <= 10）",
            ]
            for index, text in enumerate(rows, start=1):
                cells = table.add_row().cells
                cells[0].text = str(index)
                cells[1].text = text
            doc.save(docx_path)

            render_docx_to_image(docx_path, output_path, width=120, height=120)
            image = Image.open(output_path).convert("RGB")
            colors = set(image.getdata())

        self.assertIn((255, 0, 0), colors)
        self.assertIn((255, 255, 255), colors)


if __name__ == "__main__":
    unittest.main()
