from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw


NUMBER = r"[-+]?(?:\d+(?:\.\d*)?|\.\d+)(?:e[-+]?\d+)?"
LINEAR_RE = re.compile(
    rf"y\s*=\s*({NUMBER})x\s*([+-])\s*({NUMBER}).*?"
    rf"({NUMBER})\s*<=\s*x\s*<=\s*({NUMBER})",
    re.IGNORECASE,
)
VERTICAL_RE = re.compile(
    rf"x\s*=\s*({NUMBER}).*?({NUMBER})\s*<=\s*y\s*<=\s*({NUMBER})",
    re.IGNORECASE,
)
COLOR_RE = re.compile(r"颜色[:：]\s*(#[0-9a-fA-F]{6})")


Point = tuple[float, float]


@dataclass
class Segment:
    group: int
    color: str
    start: Point
    end: Point


def _close(a: Point, b: Point, tolerance: float = 1e-3) -> bool:
    return abs(a[0] - b[0]) <= tolerance and abs(a[1] - b[1]) <= tolerance


def _parse_function_text(text: str, color: str, group: int) -> Segment | None:
    linear = LINEAR_RE.search(text)
    if linear:
        m = float(linear.group(1))
        sign = -1.0 if linear.group(2) == "-" else 1.0
        b = sign * float(linear.group(3))
        x1 = float(linear.group(4))
        x2 = float(linear.group(5))
        return Segment(group=group, color=color, start=(x1, m * x1 + b), end=(x2, m * x2 + b))

    vertical = VERTICAL_RE.search(text)
    if vertical:
        x = float(vertical.group(1))
        y1 = float(vertical.group(2))
        y2 = float(vertical.group(3))
        return Segment(group=group, color=color, start=(x, y1), end=(x, y2))

    return None


def _colors_from_doc(document: Document) -> list[str]:
    colors: list[str] = []
    for paragraph in document.paragraphs:
        match = COLOR_RE.search(paragraph.text)
        if match:
            colors.append(match.group(1))
    return colors


def _segments_from_docx(docx_path: Path) -> list[Segment]:
    document = Document(docx_path)
    colors = _colors_from_doc(document)
    segments: list[Segment] = []

    for table_index, table in enumerate(document.tables):
        color = colors[table_index] if table_index < len(colors) else "#000000"
        for row in table.rows[1:]:
            if len(row.cells) < 2:
                continue
            segment = _parse_function_text(row.cells[1].text, color, table_index)
            if segment is not None:
                segments.append(segment)
    return segments


def _contours_from_segments(segments: list[Segment]) -> list[tuple[str, list[Point]]]:
    contours: list[tuple[str, list[Point]]] = []

    group_order: list[int] = []
    by_group: dict[int, list[Segment]] = {}
    for segment in segments:
        if segment.group not in by_group:
            group_order.append(segment.group)
            by_group[segment.group] = []
        by_group[segment.group].append(segment)

    for group in group_order:
        group_segments = by_group[group]
        color = group_segments[0].color
        points: list[Point] = []
        for segment in group_segments:
            if not points:
                points = [segment.start, segment.end]
                continue
            if _close(points[-1], segment.start):
                points.append(segment.end)
            elif _close(points[-1], segment.end):
                points.append(segment.start)
            else:
                if len(points) >= 3:
                    contours.append((color, points))
                points = [segment.start, segment.end]
        if len(points) >= 3:
            contours.append((color, points))
    return contours


def _hex_to_rgb(color: str) -> tuple[int, int, int]:
    return (int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16))


def _pixel_mapper(contours: list[tuple[str, list[Point]]], width: int, height: int, padding: int):
    xs = [x for _, points in contours for x, _ in points]
    ys = [y for _, points in contours for _, y in points]
    if not xs or not ys:
        raise ValueError("No function points found in DOCX")

    min_x, max_x = min(xs), max(xs)
    min_y, max_y = min(ys), max(ys)
    span_x = max(max_x - min_x, 1e-9)
    span_y = max(max_y - min_y, 1e-9)
    scale = min((width - padding * 2) / span_x, (height - padding * 2) / span_y)
    offset_x = (width - span_x * scale) / 2
    offset_y = (height - span_y * scale) / 2

    def to_pixel(point: Point) -> tuple[int, int]:
        x, y = point
        px = offset_x + (x - min_x) * scale
        py = height - (offset_y + (y - min_y) * scale)
        return (round(px), round(py))

    return to_pixel


def render_docx_to_image(
    docx_path: Path | str,
    output_path: Path | str,
    *,
    width: int = 1132,
    height: int = 1390,
    padding: int = 24,
) -> None:
    segments = _segments_from_docx(Path(docx_path))
    contours = _contours_from_segments(segments)
    to_pixel = _pixel_mapper(contours, width, height, padding)

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    for color, points in contours:
        polygon = [to_pixel(point) for point in points]
        draw.polygon(polygon, fill=_hex_to_rgb(color))
    image.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Render a Lucas function DOCX into a PNG image.")
    parser.add_argument("docx", type=Path, help="Input DOCX function list.")
    parser.add_argument("output", type=Path, help="Output PNG path.")
    parser.add_argument("--width", type=int, default=1132)
    parser.add_argument("--height", type=int, default=1390)
    parser.add_argument("--padding", type=int, default=24)
    args = parser.parse_args()

    render_docx_to_image(args.docx, args.output, width=args.width, height=args.height, padding=args.padding)
    print(args.output)


if __name__ == "__main__":
    main()
